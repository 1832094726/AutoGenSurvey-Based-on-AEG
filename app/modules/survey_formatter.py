"""
综述格式化和排版模块
支持多种输出格式：Markdown、HTML、LaTeX、PDF、Word等
"""

import logging
import re
import os
import tempfile
from typing import Dict, List, Any, Optional
from datetime import datetime
from dataclasses import dataclass
import json

@dataclass
class FormattingOptions:
    """格式化选项"""
    format_type: str  # markdown, html, latex, pdf, word
    style: str  # academic, technical, simple
    include_toc: bool = True
    include_references: bool = True
    include_figures: bool = True
    font_size: int = 12
    line_spacing: float = 1.5
    margin_size: str = "normal"  # narrow, normal, wide
    citation_style: str = "ieee"  # ieee, apa, mla

class SurveyFormatter:
    """综述格式化器"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.templates = self._load_templates()
    
    def format_survey(self, 
                     content: str, 
                     metadata: Dict[str, Any], 
                     options: FormattingOptions) -> Dict[str, Any]:
        """格式化综述"""
        try:
            self.logger.info(f"开始格式化综述，格式: {options.format_type}")
            
            # 预处理内容
            processed_content = self._preprocess_content(content, options)
            
            # 根据格式类型进行格式化
            if options.format_type == "markdown":
                formatted_content = self._format_markdown(processed_content, metadata, options)
            elif options.format_type == "html":
                formatted_content = self._format_html(processed_content, metadata, options)
            elif options.format_type == "latex":
                formatted_content = self._format_latex(processed_content, metadata, options)
            elif options.format_type == "pdf":
                formatted_content = self._format_pdf(processed_content, metadata, options)
            elif options.format_type == "word":
                formatted_content = self._format_word(processed_content, metadata, options)
            else:
                raise ValueError(f"不支持的格式类型: {options.format_type}")
            
            # 生成输出结果
            result = {
                "content": formatted_content,
                "format": options.format_type,
                "metadata": metadata,
                "formatting_options": options.__dict__,
                "generated_at": datetime.now().isoformat()
            }
            
            self.logger.info(f"综述格式化完成: {options.format_type}")
            return result
            
        except Exception as e:
            self.logger.error(f"格式化综述失败: {str(e)}")
            raise
    
    def _load_templates(self) -> Dict[str, str]:
        """加载格式化模板"""
        templates = {}
        
        # Markdown模板
        templates["markdown_academic"] = """# {title}

## 摘要
{abstract}

## 目录
{toc}

{content}

## 参考文献
{references}

---
*生成时间: {generation_time}*
"""
        
        # HTML模板
        templates["html_academic"] = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{ font-family: 'Times New Roman', serif; line-height: {line_spacing}; margin: 2cm; }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; }}
        h2 {{ color: #34495e; border-bottom: 1px solid #bdc3c7; }}
        .abstract {{ background: #f8f9fa; padding: 1em; border-left: 4px solid #3498db; }}
        .toc {{ background: #f1f2f6; padding: 1em; }}
        .references {{ font-size: 0.9em; }}
        .metadata {{ color: #7f8c8d; font-style: italic; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <div class="abstract">
        <h2>摘要</h2>
        {abstract}
    </div>
    <div class="toc">
        <h2>目录</h2>
        {toc}
    </div>
    {content}
    <div class="references">
        <h2>参考文献</h2>
        {references}
    </div>
    <div class="metadata">
        <p>生成时间: {generation_time}</p>
    </div>
</body>
</html>"""
        
        # LaTeX模板
        templates["latex_academic"] = """\\documentclass[12pt,a4paper]{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage[T1]{{fontenc}}
\\usepackage{{ctex}}
\\usepackage{{geometry}}
\\usepackage{{setspace}}
\\usepackage{{titlesec}}
\\usepackage{{hyperref}}

\\geometry{{margin=2.5cm}}
\\setstretch{{{line_spacing}}}

\\title{{{title}}}
\\author{{AutoSurvey系统}}
\\date{{{generation_time}}}

\\begin{{document}}

\\maketitle

\\begin{{abstract}}
{abstract}
\\end{{abstract}}

\\tableofcontents
\\newpage

{content}

\\begin{{thebibliography}}{{99}}
{references}
\\end{{thebibliography}}

\\end{{document}}"""
        
        return templates
    
    def _preprocess_content(self, content: str, options: FormattingOptions) -> str:
        """预处理内容"""
        # 标准化换行符
        content = content.replace('\r\n', '\n').replace('\r', '\n')
        
        # 移除多余的空行
        content = re.sub(r'\n{3,}', '\n\n', content)
        
        # 标准化标题格式
        content = self._normalize_headings(content)
        
        # 处理特殊字符
        if options.format_type in ["latex", "pdf"]:
            content = self._escape_latex_chars(content)
        elif options.format_type in ["html"]:
            content = self._escape_html_chars(content)
        
        return content.strip()
    
    def _normalize_headings(self, content: str) -> str:
        """标准化标题格式"""
        # 确保标题前后有适当的空行
        content = re.sub(r'\n*(#{1,6})\s*([^\n]+)\n*', r'\n\n\1 \2\n\n', content)
        
        # 移除标题中的多余空格
        content = re.sub(r'(#{1,6})\s+([^\n]+)', r'\1 \2', content)
        
        return content
    
    def _escape_latex_chars(self, content: str) -> str:
        """转义LaTeX特殊字符"""
        latex_chars = {
            '&': '\\&',
            '%': '\\%',
            '$': '\\$',
            '#': '\\#',
            '^': '\\textasciicircum{}',
            '_': '\\_',
            '{': '\\{',
            '}': '\\}',
            '~': '\\textasciitilde{}',
            '\\': '\\textbackslash{}'
        }
        
        for char, escape in latex_chars.items():
            content = content.replace(char, escape)
        
        return content
    
    def _escape_html_chars(self, content: str) -> str:
        """转义HTML特殊字符"""
        html_chars = {
            '&': '&amp;',
            '<': '&lt;',
            '>': '&gt;',
            '"': '&quot;',
            "'": '&#x27;'
        }
        
        for char, escape in html_chars.items():
            content = content.replace(char, escape)
        
        return content
    
    def _format_markdown(self, content: str, metadata: Dict[str, Any], options: FormattingOptions) -> str:
        """格式化为Markdown"""
        # 提取标题
        title = self._extract_title(content) or metadata.get("topic", "综述")
        
        # 生成摘要
        abstract = self._generate_abstract(content, metadata)
        
        # 生成目录
        toc = self._generate_toc_markdown(content) if options.include_toc else ""
        
        # 生成参考文献
        references = self._format_references_markdown(metadata.get("references", [])) if options.include_references else ""
        
        # 使用模板
        template = self.templates["markdown_academic"]
        
        formatted_content = template.format(
            title=title,
            abstract=abstract,
            toc=toc,
            content=content,
            references=references,
            generation_time=datetime.now().strftime("%Y年%m月%d日")
        )
        
        return formatted_content
    
    def _format_html(self, content: str, metadata: Dict[str, Any], options: FormattingOptions) -> str:
        """格式化为HTML"""
        # 提取标题
        title = self._extract_title(content) or metadata.get("topic", "综述")
        
        # 转换Markdown到HTML
        html_content = self._markdown_to_html(content)
        
        # 生成摘要
        abstract = self._generate_abstract(content, metadata)
        
        # 生成目录
        toc = self._generate_toc_html(content) if options.include_toc else ""
        
        # 生成参考文献
        references = self._format_references_html(metadata.get("references", [])) if options.include_references else ""
        
        # 使用模板
        template = self.templates["html_academic"]
        
        formatted_content = template.format(
            title=title,
            abstract=abstract,
            toc=toc,
            content=html_content,
            references=references,
            line_spacing=options.line_spacing,
            generation_time=datetime.now().strftime("%Y年%m月%d日")
        )
        
        return formatted_content
    
    def _format_latex(self, content: str, metadata: Dict[str, Any], options: FormattingOptions) -> str:
        """格式化为LaTeX"""
        # 提取标题
        title = self._extract_title(content) or metadata.get("topic", "综述")
        
        # 转换Markdown到LaTeX
        latex_content = self._markdown_to_latex(content)
        
        # 生成摘要
        abstract = self._generate_abstract(content, metadata)
        
        # 生成参考文献
        references = self._format_references_latex(metadata.get("references", [])) if options.include_references else ""
        
        # 使用模板
        template = self.templates["latex_academic"]
        
        formatted_content = template.format(
            title=title,
            abstract=abstract,
            content=latex_content,
            references=references,
            line_spacing=options.line_spacing,
            generation_time=datetime.now().strftime("%Y年%m月%d日")
        )
        
        return formatted_content
    
    def _format_pdf(self, content: str, metadata: Dict[str, Any], options: FormattingOptions) -> bytes:
        """格式化为PDF"""
        # 首先生成LaTeX
        latex_content = self._format_latex(content, metadata, options)
        
        # 使用LaTeX编译为PDF
        pdf_content = self._compile_latex_to_pdf(latex_content)
        
        return pdf_content
    
    def _format_word(self, content: str, metadata: Dict[str, Any], options: FormattingOptions) -> bytes:
        """格式化为Word文档"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            
            # 添加标题
            title = self._extract_title(content) or metadata.get("topic", "综述")
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加摘要
            abstract = self._generate_abstract(content, metadata)
            if abstract:
                doc.add_heading('摘要', level=1)
                doc.add_paragraph(abstract)
            
            # 添加目录（简化版）
            if options.include_toc:
                doc.add_heading('目录', level=1)
                toc_items = self._extract_headings(content)
                for item in toc_items:
                    doc.add_paragraph(f"{item['level'] * '  '}{item['text']}")
            
            # 添加主要内容
            self._add_content_to_word_doc(doc, content)
            
            # 添加参考文献
            if options.include_references and metadata.get("references"):
                doc.add_heading('参考文献', level=1)
                for i, ref in enumerate(metadata["references"], 1):
                    ref_text = self._format_reference_text(ref, i)
                    doc.add_paragraph(ref_text)
            
            # 保存到字节流
            import io
            doc_stream = io.BytesIO()
            doc.save(doc_stream)
            doc_stream.seek(0)
            
            return doc_stream.getvalue()
            
        except ImportError:
            self.logger.error("python-docx库未安装，无法生成Word文档")
            raise Exception("需要安装python-docx库才能生成Word文档")
    
    def _extract_title(self, content: str) -> Optional[str]:
        """提取标题"""
        # 查找第一个一级标题
        match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        return match.group(1).strip() if match else None
    
    def _generate_abstract(self, content: str, metadata: Dict[str, Any]) -> str:
        """生成摘要"""
        # 查找摘要部分
        abstract_match = re.search(r'##?\s*摘要\s*\n(.*?)(?=\n##|\Z)', content, re.DOTALL | re.IGNORECASE)
        if abstract_match:
            return abstract_match.group(1).strip()
        
        # 如果没有找到摘要，生成简单摘要
        word_count = metadata.get("word_count", 0)
        algorithm_count = metadata.get("algorithm_mentions", 0)
        
        return f"本综述包含约{word_count}字，涵盖{algorithm_count}个重要算法，对相关领域进行了全面分析。"
    
    def _generate_toc_markdown(self, content: str) -> str:
        """生成Markdown格式目录"""
        headings = self._extract_headings(content)
        toc_lines = []
        
        for heading in headings:
            indent = "  " * (heading["level"] - 1)
            link = heading["text"].lower().replace(" ", "-")
            toc_lines.append(f"{indent}- [{heading['text']}](#{link})")
        
        return "\n".join(toc_lines)
    
    def _generate_toc_html(self, content: str) -> str:
        """生成HTML格式目录"""
        headings = self._extract_headings(content)
        toc_lines = ["<ul>"]
        
        for heading in headings:
            link = heading["text"].lower().replace(" ", "-")
            toc_lines.append(f"<li><a href=\"#{link}\">{heading['text']}</a></li>")
        
        toc_lines.append("</ul>")
        return "\n".join(toc_lines)
    
    def _extract_headings(self, content: str) -> List[Dict[str, Any]]:
        """提取标题"""
        headings = []
        lines = content.split('\n')
        
        for line in lines:
            match = re.match(r'^(#{1,6})\s+(.+)$', line.strip())
            if match:
                level = len(match.group(1))
                text = match.group(2).strip()
                headings.append({"level": level, "text": text})
        
        return headings
    
    def _markdown_to_html(self, content: str) -> str:
        """简单的Markdown到HTML转换"""
        # 标题转换
        content = re.sub(r'^(#{1,6})\s+(.+)$', lambda m: f'<h{len(m.group(1))}>{m.group(2)}</h{len(m.group(1))}>', content, flags=re.MULTILINE)
        
        # 粗体转换
        content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', content)
        
        # 斜体转换
        content = re.sub(r'\*(.*?)\*', r'<em>\1</em>', content)
        
        # 段落转换
        paragraphs = content.split('\n\n')
        html_paragraphs = []
        
        for para in paragraphs:
            para = para.strip()
            if para and not para.startswith('<h'):
                html_paragraphs.append(f'<p>{para}</p>')
            else:
                html_paragraphs.append(para)
        
        return '\n'.join(html_paragraphs)
    
    def _markdown_to_latex(self, content: str) -> str:
        """简单的Markdown到LaTeX转换"""
        # 标题转换
        content = re.sub(r'^#{1}\s+(.+)$', r'\\section{\1}', content, flags=re.MULTILINE)
        content = re.sub(r'^#{2}\s+(.+)$', r'\\subsection{\1}', content, flags=re.MULTILINE)
        content = re.sub(r'^#{3}\s+(.+)$', r'\\subsubsection{\1}', content, flags=re.MULTILINE)
        
        # 粗体转换
        content = re.sub(r'\*\*(.*?)\*\*', r'\\textbf{\1}', content)
        
        # 斜体转换
        content = re.sub(r'\*(.*?)\*', r'\\textit{\1}', content)
        
        return content
    
    def _format_references_markdown(self, references: List[Dict[str, Any]]) -> str:
        """格式化Markdown参考文献"""
        if not references:
            return ""
        
        ref_lines = []
        for i, ref in enumerate(references, 1):
            ref_text = self._format_reference_text(ref, i)
            ref_lines.append(f"{i}. {ref_text}")
        
        return "\n".join(ref_lines)
    
    def _format_references_html(self, references: List[Dict[str, Any]]) -> str:
        """格式化HTML参考文献"""
        if not references:
            return ""
        
        ref_lines = ["<ol>"]
        for ref in references:
            ref_text = self._format_reference_text(ref)
            ref_lines.append(f"<li>{ref_text}</li>")
        ref_lines.append("</ol>")
        
        return "\n".join(ref_lines)
    
    def _format_references_latex(self, references: List[Dict[str, Any]]) -> str:
        """格式化LaTeX参考文献"""
        if not references:
            return ""
        
        ref_lines = []
        for ref in references:
            ref_text = self._format_reference_text(ref)
            ref_lines.append(f"\\bibitem{{ref}} {ref_text}")
        
        return "\n".join(ref_lines)
    
    def _format_reference_text(self, ref: Dict[str, Any], index: Optional[int] = None) -> str:
        """格式化单个参考文献"""
        title = ref.get("title", "未知标题")
        authors = ref.get("authors", [])
        year = ref.get("year", "未知年份")
        venue = ref.get("venue", "")
        
        author_text = ", ".join(authors) if authors else "未知作者"
        
        if venue:
            return f"{author_text}. {title}. {venue}, {year}."
        else:
            return f"{author_text}. {title}. {year}."
    
    def _compile_latex_to_pdf(self, latex_content: str) -> bytes:
        """编译LaTeX为PDF"""
        try:
            import subprocess
            
            # 创建临时目录
            with tempfile.TemporaryDirectory() as temp_dir:
                tex_file = os.path.join(temp_dir, "survey.tex")
                
                # 写入LaTeX文件
                with open(tex_file, 'w', encoding='utf-8') as f:
                    f.write(latex_content)
                
                # 编译PDF
                result = subprocess.run(
                    ['pdflatex', '-interaction=nonstopmode', tex_file],
                    cwd=temp_dir,
                    capture_output=True,
                    text=True
                )
                
                pdf_file = os.path.join(temp_dir, "survey.pdf")
                
                if os.path.exists(pdf_file):
                    with open(pdf_file, 'rb') as f:
                        return f.read()
                else:
                    raise Exception(f"PDF编译失败: {result.stderr}")
                    
        except FileNotFoundError:
            self.logger.error("pdflatex未安装，无法编译PDF")
            raise Exception("需要安装LaTeX才能生成PDF")
        except Exception as e:
            self.logger.error(f"PDF编译失败: {str(e)}")
            # 返回LaTeX内容作为fallback
            return latex_content.encode('utf-8')
    
    def _add_content_to_word_doc(self, doc, content: str):
        """将内容添加到Word文档"""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 处理标题
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                doc.add_heading(text, level=min(level, 3))
            else:
                # 普通段落
                doc.add_paragraph(line)
